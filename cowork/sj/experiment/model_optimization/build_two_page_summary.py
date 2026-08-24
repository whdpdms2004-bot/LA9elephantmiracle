from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = (
    ROOT
    / "experiment"
    / "model_optimization"
    / "LG_Aimers_인사이트_성능_요약_2p_20260812_v2.docx"
)

# Preset: compact_reference_guide.
# Named override `two_page_density`: 2-page meeting brief requirement.
FONT = "Malgun Gothic"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "E9F4EE"
GOLD = "FFF4D6"
RED = "FDECEC"
BORDER = "C9D3DF"
WHITE = "FFFFFF"


def color(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def font_run(run, size=9.4, bold=False, value="222222", italic=False):
    run.font.name = FONT
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color(value)


def set_cell_margins(cell, top=75, start=120, bottom=75, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, amount in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(amount))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_table_borders(table, border_color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), border_color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    if sum(widths_dxa) != 9360:
        raise ValueError(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for tag, attrs in (
        ("tblW", {"w": "9360", "type": "dxa"}),
        ("tblInd", {"w": str(indent_dxa), "type": "dxa"}),
        ("tblLayout", {"type": "fixed"}),
    ):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        for key, value in attrs.items():
            node.set(qn(f"w:{key}"), value)
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_header(row):
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(node)


def add_page_field(paragraph):
    run = paragraph.add_run()
    font_run(run, 8, False, MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def create_numbering(doc: Document, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(text)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def configure(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = color("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h1.font.size = Pt(14.2)
    h1.font.bold = True
    h1.font.color.rgb = color(BLUE)
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    h2.font.size = Pt(11.5)
    h2.font.bold = True
    h2.font.color.rgb = color(DARK_BLUE)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    p.paragraph_format.space_after = Pt(0)
    font_run(p.add_run("LG Aimers | 제구 성공 확률 예측"), 8.2, True, MUTED)
    font_run(p.add_run("\t2026-08-12 업데이트"), 8.2, False, MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    p.paragraph_format.space_after = Pt(0)
    font_run(p.add_run("Internal working brief"), 8, False, MUTED)
    font_run(p.add_run("\t2-page brief | 2026-08-12"), 8, False, MUTED)


def title_block(doc: Document, page_two=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    font_run(
        p.add_run("PERFORMANCE & NEXT MOVE" if page_two else "MODEL & DATA BRIEF"),
        9.0,
        True,
        BLUE,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    font_run(
        p.add_run(
            "최종 구조·제출안·다음 의사결정"
            if page_two
            else "제구 성공 확률 예측: 인사이트 및 성능 요약"
        ),
        19 if page_two else 20,
        True,
        NAVY,
    )
    if not page_two:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        font_run(
            p.add_run("EDA · success prior · 투수/타자 군집 · reverse 분리 · seed 안정화"),
            9.3,
            False,
            MUTED,
        )


def callout(doc: Document, label: str, text: str, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, fill, "0")
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.08
    font_run(p.add_run(f"{label}  "), 9.1, True, DARK_BLUE)
    font_run(p.add_run(text), 9.1, False, "222222")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)


def numbered_item(doc: Document, number_id: int, lead: str, text: str):
    p = doc.add_paragraph()
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
    num_pr.get_or_add_numId().set(qn("w:val"), str(number_id))
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.06
    p.paragraph_format.keep_together = True
    font_run(p.add_run(lead), 9.1, True, DARK_BLUE)
    font_run(p.add_run(text), 9.1, False)


def bullet_item(doc: Document, number_id: int, lead: str, text: str):
    numbered_item(doc, number_id, lead, text)


def comparison_table(doc: Document, headers, rows, widths, fills=None, size=8.15):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        font_run(p.add_run(header), 8.25, True, NAVY)
    repeat_header(table.rows[0])
    for row_index, row in enumerate(rows, start=1):
        fill = fills[row_index - 1] if fills else None
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            if fill:
                shade(cell, fill)
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if column_index in range(1, len(headers) - 1)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            font_run(p.add_run(str(value)), size, column_index == 0, "222222")
    set_table_geometry(table, widths)
    set_table_borders(table)
    return table


def build() -> Path:
    doc = Document()
    configure(doc)
    doc.core_properties.title = "LG Aimers 제구 성공 확률 예측 - 인사이트 및 성능 요약"
    doc.core_properties.subject = "submit_013까지 반영한 최신 인사이트·성능 2페이지 요약"
    doc.core_properties.author = "LG Aimers Team"

    # PAGE 1 — evidence and performance evolution
    title_block(doc)
    callout(
        doc,
        "핵심 결론",
        "성공률 보정 피처가 안정적인 기반을 만들었고, 3번 실패(reverse)를 상황 평균에서 분리한 "
        "투수유형×타자유형 잔차가 추가 성능을 만들었다. 최종 3-seed 보정 앙상블의 2024 BSS는 812.704다.",
        GREEN,
    )

    doc.add_heading("1. 데이터와 검증에서 확인한 사실", level=1)
    n1 = create_numbering(doc)
    numbered_item(
        doc,
        n1,
        "시즌 drift — ",
        "전체 제구 성공률은 2019년 56.47%에서 2024년 48.61%로 하락했다. 무작위 검증보다 연도 순방향 검증과 최근 시즌 가중치가 중요하다.",
    )
    numbered_item(
        doc,
        n1,
        "3번 실패의 독립성 — ",
        "reverse 유효 표본 비율은 전체 22.90%, 2024년 24.80%다. 단순 시즌 prior는 미약했지만 볼카운트·좌우 상황 평균을 제거한 잔차는 두 검증연도에서 개선됐다.",
    )
    numbered_item(
        doc,
        n1,
        "시간 누수 통제 — ",
        "2024 검증은 2023년 이하 train/TrackMan만 사용했다. TrackMan은 투수-시즌 500구 이상만 허용했으며 최종 2025 모델에서만 2024를 포함했다.",
    )
    numbered_item(
        doc,
        n1,
        "군집 coverage — ",
        "2024 행 기준 TrackMan 프로필 coverage는 57.34%, 투수유형×타자유형 pair coverage는 72.78%였다. 미매칭은 hand별 rookie/control-only 유형으로 분리했다.",
    )
    numbered_item(
        doc,
        n1,
        "좌우·선수 매칭 — ",
        "메인 hand는 1=좌, 2=우로 확인했다. TrackMan crosswalk 336명 중 335명이 hand와 일치했으며, 1건 불일치는 ID와 hand 정보를 분리 보존했다.",
    )

    doc.add_heading("2. 성능이 오른 순서", level=1)
    feature_rows = [
        ("XGB V2R200 + TM500", "단일", "774.484", "재현 기준선"),
        ("성공률 adjusted 2개", "단일", "784.557", "피처 기반 최고"),
        ("adjusted + performance", "혼합", "801.147", "기존 최고 007"),
        ("성공 상성 보정", "혼합", "806.488", "F23/F24 개선"),
        ("reverse 독립 보정", "혼합", "810.098", "3번 분리 효과"),
        ("reverse 전용 타자 군집", "혼합", "810.257", "좌4/우6"),
        ("reverse 3-seed 평균", "혼합", "812.704", "최종 013"),
    ]
    comparison_table(
        doc,
        ["방법", "형태", "2024 BSS", "판단"],
        feature_rows,
        [4100, 1050, 1550, 2660],
        [None, GOLD, None, GREEN, GREEN, GREEN, GREEN],
        8.0,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    font_run(p.add_run("핵심 해석  "), 8.5, True, BLUE)
    font_run(
        p.add_run(
            "hard/soft 군집 ID와 matchup 피처의 XGBoost 직접 투입은 악화됐다. 군집은 희소 상성 통계의 smoothing 단위로 쓸 때 효과가 났다. Optuna 25회에서도 trial93이 최고여서 핵심은 파라미터보다 피처 설계였다."
        ),
        8.5,
        False,
        MUTED,
    )

    doc.add_page_break()

    # PAGE 2 — final architecture and submission decision
    title_block(doc, page_two=True)
    callout(
        doc,
        "현재 최고",
        "submit_013의 2024 BSS는 812.704로 기존 007(801.147)보다 +11.557 높다. "
        "F23/F24 ΔBrier도 각각 -0.00001436/-0.00003676으로 같은 방향이다. 실제 Public LB는 Cat 838.492, XGB 873.075이며 선두권 참고 점수는 약 1100이다.",
        LIGHT_BLUE,
    )

    doc.add_heading("3. 최종 모델 구조", level=1)
    n2 = create_numbering(doc)
    numbered_item(
        doc, n2, "투수 유형 — ",
        "TrackMan 물리+과거 제구 프로필을 PCA 8차원으로 축약하고 diagonal GMM으로 좌 2·우 4군집을 구성했다.",
    )
    numbered_item(
        doc, n2, "성공 상성 — ",
        "타자 손별 KMeans 좌 3·우 4, smoothing 1000, 반감기 1년, Ridge alpha 10의 correction을 0.25 반영했다.",
    )
    numbered_item(
        doc, n2, "reverse 상성 — ",
        "시즌·투수손·타자손·볼카운트 평균을 제거하고 reverse 전용 타자 KMeans 좌 4·우 6으로 유형 쌍 잔차를 만들었다.",
    )
    numbered_item(
        doc, n2, "안정화·혼합 — ",
        "reverse seed 17/2026/4099의 Ridge(alpha 1000) correction을 평균해 0.55 반영하고, 전체 제구 branch를 performance와 0.6085:0.3915로 혼합했다.",
    )

    doc.add_heading("4. 제출 후보 성능", level=1)
    submission_rows = [
        ("submit_007.zip", "기존 adjusted", "801.147", "비교 기준"),
        ("submit_009.zip", "성공 상성 안정형", "805.564", "reverse 제외"),
        ("submit_011.zip", "성공+reverse", "810.098", "reverse 재현형"),
        ("submit_012.zip", "reverse 전용 군집", "810.257", "이전 최고"),
        ("submit_013.zip", "reverse 3-seed", "812.704", "최종 최우선"),
    ]
    comparison_table(
        doc,
        ["파일", "핵심 설계", "2024 BSS", "역할"],
        submission_rows,
        [2200, 3300, 1600, 2260],
        [None, GOLD, GREEN, GREEN, GREEN],
        8.0,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    font_run(p.add_run("검증 완료  "), 8.45, True, BLUE)
    font_run(
        p.add_run(
            "013은 2024까지 전체 재학습했으며 상대경로, CRC, 모델 파일, 30자 미만 파일명, 245,789행 추론(16.48초), 확률 검사를 통과했다."
        ),
        8.45,
        False,
        MUTED,
    )

    doc.add_heading("5. 제출·후속 의사결정", level=1)
    n3 = create_numbering(doc)
    numbered_item(doc, n3, "첫 제출 — ", "submit_013을 먼저 제출해 기존 Public LB 873.075와 비교한다.")
    numbered_item(
        doc,
        n3,
        "두 번째 슬롯 — ",
        "013이 상승하면 012로 reverse 신호를 재확인하고, 하락하면 009로 reverse 도메인 이동을 분리한다.",
    )
    numbered_item(
        doc,
        n3,
        "상승 후 확장 — ",
        "reverse count-bucket interaction과 pair-table ensemble만 추가하고, F23/F24 동시 개선 및 TrackMan 500구 규칙을 계속 고정한다.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
